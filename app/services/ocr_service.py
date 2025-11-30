"""
OCR Service - 阿里云通义千问 VL-OCR

提供文档 OCR 识别功能，支持：
- PDF 文档（包括扫描件/图片PDF）
- 图片文件（jpg/png/webp）
- Word 文档（.docx）
- Excel 文件（.xlsx）

所有文件统一转换为图片后送入 qwen-vl-ocr 模型处理。
"""

import base64
import logging
import tempfile
import os
from typing import List, Optional
from pathlib import Path
from openai import OpenAI

from app.core.config import settings

logger = logging.getLogger(__name__)


class OCRService:
    """阿里云通义千问 VL-OCR 服务封装"""

    def __init__(self):
        self.client = OpenAI(
            api_key=settings.DASHSCOPE_API_KEY,
            base_url="https://dashscope.aliyuncs.com/compatible-mode/v1"
        )
        self.model = settings.QWEN_OCR_MODEL

    async def ocr_image(self, image_path: str) -> str:
        """
        对单张图片进行 OCR 识别

        Args:
            image_path: 图片文件路径

        Returns:
            识别出的文本内容
        """
        # 读取并编码图片
        with open(image_path, "rb") as f:
            image_data = base64.b64encode(f.read()).decode("utf-8")

        # 获取 MIME 类型
        ext = Path(image_path).suffix.lower()
        mime_map = {
            ".jpg": "image/jpeg",
            ".jpeg": "image/jpeg",
            ".png": "image/png",
            ".webp": "image/webp",
            ".gif": "image/gif",
            ".bmp": "image/bmp",
        }
        mime_type = mime_map.get(ext, "image/jpeg")

        # 调用 qwen-vl-ocr
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{image_data}"
                            },
                        },
                        {
                            "type": "text",
                            "text": "Read all the text in the image."
                        },
                    ],
                }
            ],
        )

        return response.choices[0].message.content or ""

    async def ocr_pdf(self, pdf_path: str) -> str:
        """
        对 PDF 文件进行 OCR 识别

        将 PDF 每页转换为图片，然后逐页识别

        Args:
            pdf_path: PDF 文件路径

        Returns:
            识别出的全部文本内容
        """
        try:
            import pymupdf
        except ImportError:
            raise RuntimeError("请安装 pymupdf: pip install pymupdf")

        results = []
        doc = pymupdf.open(pdf_path)

        with tempfile.TemporaryDirectory() as temp_dir:
            for page_num, page in enumerate(doc):
                # 先尝试提取文本（对于可复制文本的PDF）
                text = page.get_text().strip()

                if len(text) > 100:
                    # 文本足够多，直接使用
                    results.append(f"--- 第 {page_num + 1} 页 ---\n{text}")
                else:
                    # 文本太少（可能是扫描件），转换为图片进行 OCR
                    pix = page.get_pixmap(dpi=150)
                    img_path = os.path.join(temp_dir, f"page_{page_num + 1}.png")
                    pix.save(img_path)

                    try:
                        ocr_text = await self.ocr_image(img_path)
                        results.append(f"--- 第 {page_num + 1} 页 ---\n{ocr_text}")
                    except Exception as e:
                        logger.error(f"OCR 第 {page_num + 1} 页失败: {e}")
                        results.append(f"--- 第 {page_num + 1} 页 ---\n[OCR 失败: {str(e)}]")

        doc.close()
        return "\n\n".join(results)

    async def ocr_docx(self, docx_path: str) -> str:
        """
        对 Word 文档进行文本提取

        Args:
            docx_path: Word 文档路径

        Returns:
            提取出的文本内容
        """
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("请安装 python-docx: pip install python-docx")

        doc = Document(docx_path)
        results = []

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                results.append(para.text)

        # 提取表格
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                results.append("\n".join(table_text))

        return "\n\n".join(results)

    async def ocr_xlsx(self, xlsx_path: str) -> str:
        """
        对 Excel 文件进行文本提取

        Args:
            xlsx_path: Excel 文件路径

        Returns:
            转换为 Markdown 表格的内容
        """
        try:
            import openpyxl
        except ImportError:
            raise RuntimeError("请安装 openpyxl: pip install openpyxl")

        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        results = []

        for sheet in wb.worksheets:
            sheet_content = [f"## Sheet: {sheet.title}"]

            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue

            # 构建 Markdown 表格
            for i, row in enumerate(rows):
                cells = [str(c) if c is not None else "" for c in row]
                sheet_content.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    # 添加表头分隔线
                    sheet_content.append("| " + " | ".join(["---"] * len(cells)) + " |")

            results.append("\n".join(sheet_content))

        return "\n\n".join(results)

    async def ocr_file(self, file_path: str) -> str:
        """
        智能识别文件类型并进行 OCR

        Args:
            file_path: 文件路径

        Returns:
            识别/提取出的文本内容
        """
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            return await self.ocr_pdf(file_path)
        elif ext in [".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp"]:
            return await self.ocr_image(file_path)
        elif ext == ".docx":
            return await self.ocr_docx(file_path)
        elif ext == ".xlsx":
            return await self.ocr_xlsx(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")


# 全局单例
_ocr_service: Optional[OCRService] = None


def get_ocr_service() -> OCRService:
    """获取 OCR 服务实例"""
    global _ocr_service
    if _ocr_service is None:
        _ocr_service = OCRService()
    return _ocr_service
